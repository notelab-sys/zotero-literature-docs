# -*- coding: utf-8 -*-
"""Generate a literature review (Word + PDF) from review_content.json.

Layout follows the style of Acta Horticulturae Sinica (园艺学报):
- In-text citations are author-year style: （Zhang et al.，2015）.
- The reference list is alphabetical by first author.
- PDF uses SimSun (body) / SimHei (title & level-2 headings) with the
  journal's sizes and line spacing; headings are numbered (1, 1.1, ...).
- Optional tables and figures are rendered with the journal's caption style
  (Chinese caption + English caption; table captions above, figure captions
  below).

Content JSON schema (same as before, plus optional keys):
  "keywords": "A；B；C"
  groups may contain:
    "table": {"caption": "表1 ...", "caption_en": "Table 1 ...",
              "columns": [...], "rows": [[...], ...], "widths": [mm, ...] (optional)}
    "image": {"path": "...", "width": 120 (mm, optional),
              "caption": "图1 ...", "caption_en": "Fig. 1 ..."}
"""

import json
import re
import struct
import sys
import unicodedata
import zipfile
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:  # noqa: BLE001
    pass

WORK = Path(__file__).parent
sys.path.insert(0, str(WORK))
sys.path.insert(0, str(WORK / "pylibs"))
sys.path.insert(0, str(WORK / "pylibs2"))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402
from fpdf import FPDF  # noqa: E402

CONTENT_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else WORK / "review_content.json"
CONTENT = json.loads(CONTENT_PATH.read_text(encoding="utf-8"))
EN = str(CONTENT.get("lang", "")).lower() == "en"
DATA_PATH = WORK / "data.json"
for a in sys.argv[1:]:
    if a.startswith("--data"):
        i = sys.argv.index(a)
        if i + 1 < len(sys.argv):
            DATA_PATH = Path(sys.argv[i + 1])
DATA = json.loads(DATA_PATH.read_text(encoding="utf-8"))
OUT_DIR = WORK.parent / "outputs"
OUT_DIR.mkdir(exist_ok=True)
OUT_NAME = CONTENT.get("output", "文献综合评述")

DARK = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)

CITE_RE = re.compile(r"\[@([A-Z0-9]{8})\]")
CITE_RUN = re.compile(r"(?:\[@[A-Z0-9]{8}\]\s*)+")
CN_NUM = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
          "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}

FONT_DIR = Path(r"C:\Windows\Fonts")
SIM_SUN = FONT_DIR / "simsun.ttc"
SIM_HEI = FONT_DIR / "simhei.ttf"

# Gene symbols / gene names that should be italic in Word documents.
ITALIC_WORDS = {
    "AP1", "AP2", "AP3", "PI", "AG", "SEP1", "SEP2", "SEP3", "SEP4", "FUL", "STK",
    "LFY", "LFY1", "LFY2", "WUS", "CLV1", "PTL", "CYC", "CYC2a", "CYC2b", "DICH",
    "RAD", "DIV", "DRIF", "CHS", "CHI", "F3H", "F3'H", "F3'5'H", "DFR", "ANS",
    "UFGT", "LDOX", "LAR", "ANR", "FLS", "FNS", "TT2", "TT8", "TT12", "PAP1",
    "PAP2", "SVP", "AGL6", "AGL24", "SOC1", "OAP3-1", "OAP3-2", "OAGL6-1",
    "OAGL6-2", "OPI", "OMADS2", "OMADS3", "OMADS4", "OMADS6", "OMADS7", "OMADS10",
    "OMADS11", "CsTT8", "CsDFR", "CsDFRa", "CsDFRc", "CsDFRb1", "CsDFRb3",
    "MdMYB10", "MdHY5", "MdBBX22", "MdMYB9", "MdMYB11", "MdMYB12", "mdm-miR858",
    "miR858", "FaMYB5", "FaMYB10", "FaEGL3", "FaLWD1", "FaLWD1-like", "VmMYBA1",
    "VmMYBA2", "VmMYBPA1.1", "VmMYBPA2.2", "PeMYB4L", "PeMYC4", "PeUFGT3",
    "PeCHS", "PpMYB1", "PpbHLH1", "ASR1", "ASR2", "ASR3", "ASR1-3", "RUBY1",
    "PrMYB5", "PrDFR", "PrMYBa1", "PrMYBa2", "PrMYBa3", "PrF3H", "PrANS",
    "ArMYB89", "RrMYB5", "RrMYB10", "RrDFR", "RrANR", "GmMYB4", "GmMAPK6",
    "GmIFS2", "GmANS3", "DhMYB22", "DhMYB60", "NtMYB12", "NtFLS", "NtLAR",
    "NtDFR", "NtbHLH1", "NtWD40-1", "FhFLS1", "FhFLS2", "FhMYBF", "FhMYB21L2",
    "GaYP", "GaFLS", "BnaFLS1", "BnaFLS1-1", "BnaFLS1-2", "HvDFR", "VcF3'5'H",
    "VcF3'5'H4", "RhARF8", "RhARF1", "RhCHSa", "RhCHSc", "AcMADS68", "AcMYBF110",
    "AcMYB123", "AcbHLH1", "PsMYB44", "CmMYB3-like", "NEGAN", "RTO", "CyMYB1",
    "CybHLH1", "CybHLH2", "ZmC1", "ZmR", "AN1", "AN2", "AN11", "MYB27", "MYBx",
    "DPL", "PHZ", "JAF13", "GL3", "EGL3", "TTG1", "LcMYB1", "LcbHLH1", "LcbHLH3",
    "MdbHLH3", "MdbHLH33", "GtMYB3", "GtbHLH1", "LhMYB6", "LhMYB12", "LhbHLH2",
    "OgMYB1", "PeMYB2", "PeMYB11", "PeMYB12", "HORT1", "EjCAL", "LoSVP", "SlCMB1",
    "CsMADS5", "CsMADS6", "CitAN1", "CitPH4", "NOEMI", "CsPH4", "PhDC", "ClPHT4;2",
    "ClLCYB", "SlAN2", "SlBBX20", "SlHY5", "SlAREB1", "MdNAC42", "MdMYB1",
    "LEAFY", "CYCLOIDEA", "DICHOTOMA", "DIVARICATA", "RADIALIS", "APETALA1",
    "APETALA3", "PISTILLATA", "AGAMOUS", "SEPALLATA3", "WUSCHEL", "CLAVATA1",
    "PETAL LOSS", "FLOWERING LOCUS T", "CONSTANS", "SHORT VEGETATIVE PHASE",
    "AGAMOUS-LIKE 24", "BBM", "LEC1", "LEC2", "FUS3", "ABI3", "LAFL", "ARF",
    "ARF1", "ARF2", "ARF6", "ARF8", "ARF10", "ARF16", "ARF17", "LCR", "miR394",
    "BABY BOOM", "LEAFY COTYLEDON1", "LEAFY COTYLEDON2", "FUSCA3",
    "ABSCISIC ACID INSENSITIVE3", "LEAF CURLING RESPONSIVENESS",
    "AUXIN RESPONSE FACTOR",
}
_BINOMIAL = r"\b[A-Z][a-z]{2,}(?:\s*[×x]\s*|\s+)[a-z]{3,}\b"
_GENE_ALT = "|".join(r"\b" + re.escape(w) + r"\b"
                     for w in sorted(ITALIC_WORDS, key=len, reverse=True))
_GENE_RE = re.compile(_GENE_ALT)
_PROT_CONN = r"[\s（(，,:：、与和及兼具活性]*"
STOP_SPECIES = {
    "flower", "species", "plant", "plants", "gene", "genes", "hybrid",
    "variety", "group", "family", "class", "model", "review", "research",
    "analysis", "study", "studies", "development", "regulation", "protein",
    "proteins", "pathway", "complex", "biosynthesis", "fruit", "color",
    "colour", "flesh", "box", "domain", "genes", "floral", "homeotic",
    "meristem", "flowers", "organs", "organ", "seeds", "mutants", "varieties",
    "cultivars", "transcription", "factor", "factors", "expression",
    "embryogenesis", "embryonic", "embryo", "embryos", "somatic", "regulates",
    "role", "roles", "network", "networks", "induction", "induced", "encoding",
    "encodes", "receptor", "kinase", "synthase", "reductase", "transferase",
    "hydroxylase", "oxidase", "peroxidase", "activator", "repressor",
    "regulator", "regulators", "mutant", "callus", "culture", "cultures",
    "cells", "cell", "tissue", "tissues", "maturation", "morphogenesis",
    "organogenesis", "regeneration", "proliferation", "during", "using",
    "through", "with", "into", "from", "for", "the", "of", "in", "on", "at",
    "by", "to", "and", "response", "responses", "signal", "signaling",
    "signalling", "stress", "resistance", "tolerance", "activity",
    "activities", "function", "functions",
}
PLANT_GENERA = {
    "Arabidopsis", "Antirrhinum", "Mimulus", "Glycine", "Fragaria", "Vaccinium",
    "Camellia", "Paeonia", "Dendrobium", "Phalaenopsis", "Oncidium",
    "Paphiopedilum", "Lilium", "Chimonanthus", "Hosta", "Nicotiana", "Solanum",
    "Vitis", "Oryza", "Triticum", "Zea", "Gossypium", "Brassica", "Citrus",
    "Malus", "Rosa", "Acer", "Nelumbo", "Nitraria", "Ipomoea", "Euryale",
    "Curcuma", "Rhododendron", "Prunus", "Freesia", "Narcissus", "Musa",
    "Escherichia", "Actinidia", "Cymbidium", "Iris", "Osmanthus", "Anthurium",
    "Petunia", "Gentiana", "Delphinium", "Senecio", "Clitoria", "Osteospermum",
    "Phaseolus", "Capsicum", "Ginkgo", "Ribes", "Hordeum", "Pyrus", "Mangifera",
    "Chirita", "Pisum", "Medicago", "Lotus", "Helianthus", "Aquilegia",
    "Nigella", "Trollius", "Coffea", "Pinus", "Abies", "Picea", "Magnolia",
    "Populus", "Eucalyptus", "Quercus", "Betula", "Alnus", "Salix", "Ulmus",
    "Fraxinus", "Tilia", "Platanus", "Robinia", "Acacia", "Carica", "Olea",
    "Elaeis", "Cocos", "Phoenix", "Theobroma", "Rubus", "Sorghum", "Setaria",
    "Brachypodium", "Saccharum", "Chrysanthemum", "Gerbera", "Tulipa",
    "Gladiolus", "Agapanthus", "Hippeastrum", "Lycoris", "Clivia",
    "Zantedeschia", "Spathiphyllum", "Caladium", "Alocasia", "Colocasia",
    "Monstera", "Philodendron", "Epipremnum", "Zingiber", "Hedychium",
    "Alpinia", "Amomum", "Elettaria", "Canna", "Heliconia", "Strelitzia",
    "Ravenala", "Juglans", "Castanea", "Fagus", "Corylus", "Morus", "Ficus",
    "Cannabis", "Humulus", "Beta", "Spinacia", "Cucumis", "Citrullus",
    "Cucurbita", "Momordica", "Lagenaria", "Daucus", "Apium", "Petroselinum",
    "Coriandrum", "Foeniculum", "Allium", "Asparagus", "Physalis", "Withania",
    "Atropa", "Datura", "Brugmansia", "Cestrum", "Lycium", "Ilex", "Moringa",
    "Hibiscus", "Abelmoschus", "Durio", "Ceiba", "Adansonia", "Raphanus",
    "Sinapis", "Eutrema", "Camelina", "Thlaspi", "Capsella", "Arabis",
    "Cardamine", "Nasturtium", "Lepidium", "Matthiola", "Erysimum", "Barbarea",
    "Armoracia", "Cochlearia", "Draba", "Lactuca", "Cichorium", "Taraxacum",
    "Artemisia", "Dendranthema", "Calendula", "Tagetes", "Zinnia", "Dahlia",
    "Arnica", "Echinacea", "Rudbeckia", "Solidago", "Aster", "Bellis",
    "Erigeron", "Tanacetum", "Achillea", "Matricaria", "Bidens", "Cosmos",
    "Coreopsis", "Gaillardia", "Helenium", "Liatris", "Centaurea", "Cynara",
    "Silybum", "Carduus", "Cirsium", "Onopordum", "Echinops", "Arctium",
    "Carthamus", "Saussurea",
}


# --------------------------------------------------------------------------
# Citation / reference helpers (author-year, 园艺学报 style)
# --------------------------------------------------------------------------

def clean_title(title):
    title = re.sub(r"<[^>]+>", "", title or "")
    return title.replace("&amp;", "&").strip()


def surname(name):
    parts = str(name).strip().split()
    return parts[-1] if parts else ""


def initials_name(name):
    parts = str(name).strip().split()
    if len(parts) <= 1:
        n = str(name).strip()
        return n[:1].upper() + n[1:] if n else n
    initials = " ".join(p[0].upper() for p in parts[:-1] if p and p[0].isalpha())
    surname = parts[-1]
    surname = surname[:1].upper() + surname[1:]
    return f"{surname} {initials}".strip()


def author_list(author_str):
    return [a.strip() for a in str(author_str or "").split(",") if a.strip()]


def cite_str(key):
    e = next((x for x in DATA if x["key"] == key), None)
    if not e:
        return str(key)
    authors = author_list(e.get("author", ""))
    year = e.get("year", "")
    sep = ", " if EN else "，"
    if not authors:
        return (f"Anonymous, {year}" if year else "Anonymous") if EN else (f"佚名，{year}" if year else "佚名")
    first = surname(authors[0])
    if len(authors) == 1:
        s = first
    elif len(authors) == 2:
        s = f"{first} & {surname(authors[1])}"
    else:
        s = f"{first} et al."
    return f"{s}{sep}{year}" if year else s


def render_text(text, by_key=None):
    def repl(m):
        keys = CITE_RE.findall(m.group(0))
        if EN:
            return "(" + "; ".join(cite_str(k) for k in keys) + ")"
        return "（" + "；".join(cite_str(k) for k in keys) + "）"
    return CITE_RUN.sub(repl, text)


def pdf_safe(text):
    """Strip Latin diacritics for fonts without those glyphs (e.g. SimSun),
    while leaving Chinese punctuation untouched."""
    out = []
    for ch in text:
        o = ord(ch)
        if 0x00C0 <= o <= 0x024F or 0x1E00 <= o <= 0x1EFF:
            d = unicodedata.normalize("NFD", ch)
            out.append(d[0] if d else ch)
        else:
            out.append(ch)
    return "".join(out)


def ref_entry(key):
    return next((x for x in DATA if x["key"] == key), None)


def ref_sort_key(key):
    e = ref_entry(key) or {}
    authors = author_list(e.get("author", ""))
    return (surname(authors[0]) if authors else "").lower()


def ref_text(key):
    e = ref_entry(key)
    if not e:
        return ""
    authors = [initials_name(a) for a in author_list(e.get("author", ""))]
    head = ", ".join(authors)
    year = str(e.get("year", ""))
    line = head + "."
    if year:
        line += " " + year
    line += ". " + sentence_case_title(clean_title(e.get("title", ""))) + "."
    journal = e.get("journal", "")
    if journal:
        loc = journal
        vol = e.get("volume", "")
        num = e.get("number", "")
        pages = e.get("pages", "")
        if vol:
            loc += f", {vol}"
            if num:
                loc += f"({num})"
            if pages:
                loc += f": {pages.replace('-', '–')}"
        line += f" {loc}."
    doi = e.get("doi", "")
    if doi:
        line += f" DOI: {doi}."
    return line


def sentence_case_title(title):
    """Convert a reference title to sentence case (first word capitalized,
    middle words lowercase) while preserving gene names, Latin binomials and
    naming authorities (e.g. 'Linden.', 'Pierre ex A. Froehner')."""
    spans = italic_spans(title)
    first_done = False
    auth_budget = 0

    def is_binomial(seg):
        return bool(re.match(r"^([A-Z][a-z]{2,})[\s×]+[a-z]{3,}", seg))

    def scase(seg):
        nonlocal first_done, auth_budget
        words = re.findall(r"[A-Za-z][A-Za-z0-9'./\-]*|[^A-Za-z]+", seg)
        out = []
        for w in words:
            if not re.search(r"[A-Za-z]", w):
                out.append(w)
                continue
            m = re.match(r"^([A-Za-z0-9'./\-]+)(.*)$", w)
            core, tail = m.group(1), m.group(2)
            if len(core) > 1 and core.isupper():
                # Keep acronyms / all-caps gene abbreviations uppercase.
                out.append(core + tail)
                first_done = True
                continue
            if core in PLANT_GENERA:
                out.append(core[:1].upper() + core[1:].lower() + tail)
                first_done = True
                continue
            if auth_budget > 0 and core[:1].isupper():
                # Naming authority: keep first letter capitalized, regular type.
                out.append(core[:1].upper() + core[1:] + tail)
                auth_budget -= 1
                first_done = True
                continue
            if not first_done:
                out.append(core[:1].upper() + core[1:].lower() + tail)
                first_done = True
            else:
                out.append(core.lower() + tail)
        return "".join(out)

    parts = []
    pos = 0
    for s, e in spans:
        parts.append(scase(title[pos:s]))
        parts.append(title[s:e])  # gene names / binomials keep their case
        if re.search(r"[A-Za-z]", title[s:e]):
            first_done = True
            if is_binomial(title[s:e]):
                auth_budget = 4
        pos = e
    parts.append(scase(title[pos:]))
    return "".join(parts)


def collect_texts(sections):
    for sec in sections:
        for group in sec.get("groups", []):
            for item in group.get("items", []):
                yield item


def build_citations(sections):
    order = []
    for text in collect_texts(sections):
        for m in CITE_RE.finditer(text):
            k = m.group(1)
            if k not in order:
                order.append(k)
    return order


def validate_citations(order):
    data_keys = {e["key"] for e in DATA}
    missing = [k for k in order if k not in data_keys]
    if missing:
        raise SystemExit(
            "ERROR: citations missing from data.json: " + ", ".join(sorted(missing))
            + "\nRun fetch_items.py first or fix the [@KEY] tokens in the content file."
        )
    cited = set(order)
    uncited = sorted(data_keys - cited)
    if uncited:
        print(f"NOTE: {len(uncited)} item(s) in data.json are not cited "
              f"(they will not appear in the reference list): {', '.join(uncited[:10])}")


def section_number(title):
    m = re.match(r"^([一二三四五六七八九十]+)[、．.]\s*(.*)", title)
    if m:
        return CN_NUM.get(m.group(1), 0), m.group(2)
    return None, title


def sub_number(title):
    m = re.match(r"^(\d+)[、．.]\s*(.*)", title)
    if m:
        return int(m.group(1)), m.group(2)
    return None, title


def set_east_asia(run, name="SimSun"):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, size=11, color=None, bold=False, align=None, after=6,
             indent=False, east="SimSun", italics=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.3
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if italics:
        add_italicized_runs(p, text, size=size, color=color, bold=bold, east=east)
        return p
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.bold = bold
    set_east_asia(r, east)
    if EN:
        r.font.name = "Times New Roman"
    return p


def _add_run(p, text, size, color, bold, east, italic):
    if not text:
        return
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    set_east_asia(r, east)
    if EN:
        r.font.name = "Times New Roman"


def add_italicized_runs(p, text, size=10.5, color=None, bold=False, east="SimSun"):
    """Split text into runs, italicizing Latin binomials and gene symbols."""
    spans = italic_spans(text)
    merged = []
    for s, e in spans:
        if merged and s <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], e))
        else:
            merged.append((s, e))
    pos = 0
    for s, e in merged:
        if s > pos:
            _add_run(p, text[pos:s], size, color, bold, east, False)
        _add_run(p, text[s:e], size, color, bold, east, True)
        pos = e
    if pos < len(text):
        _add_run(p, text[pos:], size, color, bold, east, False)


def italic_spans(text):
    """Return (start, end) spans that should be italic."""
    spans = []
    for m in re.finditer(_BINOMIAL, text):
        seg = m.group(0)
        mm = re.match(r"^([A-Z][a-z]{2,})[\s×]+(.+)$", seg, re.S)
        if mm and mm.group(1) in PLANT_GENERA:
            if mm.group(2).lower() in STOP_SPECIES:
                spans.append((m.start(), m.start() + len(mm.group(1))))
            else:
                spans.append((m.start(), m.end()))
    for m in _GENE_RE.finditer(text):
        if not _protein_or_enzyme(text, m.start(), m.end()):
            spans.append((m.start(), m.end()))
    return spans


def _protein_or_enzyme(text, s, e):
    """True when the token denotes a protein/enzyme abbreviation (regular type)."""
    before = text[max(0, s - 16):s]
    after = text[e:e + 14]
    if re.search(r"(?:蛋白|酶)" + _PROT_CONN + r"$", before):
        return True
    if re.search(r"^" + _PROT_CONN + r"(?:蛋白|酶)", after):
        return True
    if re.search(r"^[（(][^）)]{0,14}(?:蛋白|酶)", after):
        return True
    return False


def add_page_number(doc):
    """Add centered footer page number: 第 X 页 / 共 Y 页."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def field(instr):
        r = p.add_run()
        r.font.size = Pt(9)
        set_east_asia(r)
        f1 = OxmlElement("w:fldChar")
        f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        f2 = OxmlElement("w:fldChar")
        f2.set(qn("w:fldCharType"), "end")
        r._r.append(f1)
        r._r.append(it)
        r._r.append(f2)

    def txt(text):
        r = p.add_run(text)
        r.font.size = Pt(9)
        set_east_asia(r)

    if EN:
        txt("Page ")
        field("PAGE")
        txt(" of ")
        field("NUMPAGES")
    else:
        txt("第 ")
        field("PAGE")
        txt(" 页 / 共 ")
        field("NUMPAGES")
        txt(" 页")


# --------------------------------------------------------------------------
# Word document
# --------------------------------------------------------------------------

def add_docx_table(doc, spec):
    caption = spec.get("caption", "")
    caption_en = spec.get("caption_en", "")
    columns = spec.get("columns", [])
    rows = spec.get("rows", [])
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption)
        r.font.size = Pt(10.5)
        set_east_asia(r, "SimHei")
    if caption_en:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption_en)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        set_east_asia(r)
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _docx_three_line(table)
    for j, c in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(str(c))
        r.font.size = Pt(9)
        set_east_asia(r, "SimHei")
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            add_italicized_runs(para, render_text(str(val)), size=9)
    # Keep the table on one page when it fits (same as PDF); Word will split
    # a table that is taller than a page, with rows never breaking mid-row.
    n_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        if i == n_rows - 1:
            continue
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _docx_three_line(table):
    """Apply booktabs-style borders: top/bottom thick, header underline only."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for child in list(borders):
        borders.remove(child)
    for edge, sz in [("top", "12"), ("bottom", "12")]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ["left", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = tcPr.find(qn("w:tcBorders"))
        if tcB is None:
            tcB = OxmlElement("w:tcBorders")
            tcPr.append(tcB)
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        tcB.append(b)
    # Prevent rows from splitting across pages.
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit")
        trPr.append(cs)


def add_docx_figure(doc, spec, base):
    path = Path(spec["path"])
    if not path.is_absolute():
        path = base / path
    width = spec.get("width", 13)
    if path.exists():
        doc.add_picture(str(path), width=Cm(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = spec.get("caption", "")
    cap_en = spec.get("caption_en", "")
    if cap:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cap)
        r.font.size = Pt(10.5)
        set_east_asia(r, "SimHei")
    if cap_en:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cap_en)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
        set_east_asia(r)


def build_docx(ref_keys):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.line_spacing = 1.3

    sec0 = doc.sections[0]
    title_max_mm = (sec0.page_width - sec0.left_margin - sec0.right_margin) / 36000
    title_lines = docx_title_lines(CONTENT["title"], 20, title_max_mm)
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(3)
    for i, ln in enumerate(title_lines):
        r = tp.add_run(ln)
        r.font.size = Pt(20)
        r.font.name = "SimHei"
        set_east_asia(r, "SimHei")
        if i < len(title_lines) - 1:
            r.add_break()
    add_para(doc, CONTENT["subtitle"], size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    if CONTENT.get("intro"):
        add_para(doc, ("Abstract: " if EN else "摘  要：") + render_text(CONTENT["intro"]), size=10.5, after=6, italics=True)
    if CONTENT.get("keywords"):
        add_para(doc, ("Keywords: " if EN else "关键词：") + CONTENT["keywords"], size=9, color=GRAY, after=8)
    if not EN and CONTENT.get("title_en"):
        add_para(doc, CONTENT["title_en"], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    if not EN and CONTENT.get("abstract_en"):
        p = add_para(doc, "Abstract：" + CONTENT["abstract_en"], size=9, color=GRAY,
                     after=6, italics=True)
        p.paragraph_format.first_line_indent = Pt(0)
    if not EN and CONTENT.get("keywords_en"):
        add_para(doc, "Keywords：" + CONTENT["keywords_en"], size=9, color=GRAY, after=8)

    sec_idx = 0
    for sec in CONTENT["sections"]:
        n, text = section_number(sec["heading"])
        sec_idx = n if n else sec_idx + 1
        add_para(doc, f"{sec_idx}  {text}", size=14, after=5, east="SimSun")
        doc.paragraphs[-1].paragraph_format.space_before = Pt(10)
        sub_idx = 0
        for group in sec.get("groups", []):
            if group.get("table"):
                add_docx_table(doc, group["table"])
            if group.get("image"):
                add_docx_figure(doc, group["image"], WORK.parent)
            if group.get("sub"):
                sn, stext = sub_number(group["sub"])
                sub_idx = sn if sn else sub_idx + 1
                add_para(doc, f"{sec_idx}.{sub_idx}  {stext}", size=11, after=3, east="SimHei")
            for item in group.get("items", []):
                add_para(doc, render_text(item), size=10.5, after=3, indent=True, italics=True)

    add_para(doc, "References" if EN else "参考文献", size=14, after=5, east="SimSun")
    doc.paragraphs[-1].paragraph_format.space_before = Pt(12)
    for key in sorted(ref_keys, key=ref_sort_key):
        txt = ref_text(key)
        if not txt:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.first_line_indent = Pt(-16)
        add_italicized_runs(p, txt, size=8)

    add_page_number(doc)
    out = OUT_DIR / (OUT_NAME + ".docx")
    doc.save(out)
    print(f"wrote {out} ({out.stat().st_size} bytes)")
    return out


# --------------------------------------------------------------------------
# PDF (园艺学报 style)
# --------------------------------------------------------------------------

def png_size(path):
    with open(path, "rb") as fh:
        head = fh.read(24)
    w, h = struct.unpack(">II", head[16:24])
    return w, h


class ReviewPDF(FPDF):
    def footer(self):
        self.set_y(-13)
        self.set_font("TimesR" if EN else "Song", size=7.5)
        self.set_text_color(120, 125, 130)
        if EN:
            self.cell(0, 8, f"Page {self.page_no()} of {{nb}}", align="C")
        else:
            self.cell(0, 8, f"第 {self.page_no()} 页 / 共 {{nb}} 页", align="C")


def wrap_text(pdf, text, width_mm):
    lines = []
    for para in str(text).split("\n"):
        cur = ""
        for ch in para:
            if cur and pdf.get_string_width(cur + ch) > width_mm:
                lines.append(cur)
                cur = ch
            else:
                cur += ch
        lines.append(cur)
    return lines


def wrap_title(pdf, text, width_mm):
    """Character-greedy wrap so short tokens (e.g. 'MBW') are never left on
    their own line; trailing characters move to a centered second line."""
    lines = []
    cur = ""
    for ch in text:
        if cur and pdf.get_string_width(cur + ch) > width_mm:
            lines.append(cur)
            cur = ch
        else:
            cur += ch
    if cur:
        lines.append(cur)
    return lines


def draw_mixed(pdf, x, y, width, lh, text, size,
               align="L", reg_font="Song", italic_font="TimesI"):
    """Draw one PDF line with italic Latin/gene segments in Times Italic."""
    if EN:
        reg_font = "TimesR"
    spans = italic_spans(text)
    parts = []
    pos = 0
    for s, e in spans:
        if s > pos:
            parts.append((text[pos:s], False))
        parts.append((text[s:e], True))
        pos = e
    if pos < len(text):
        parts.append((text[pos:], False))

    def seg_w(seg, it):
        pdf.set_font(italic_font if it else reg_font, size=size)
        return pdf.get_string_width(seg)

    total = sum(seg_w(seg, it) for seg, it in parts)
    x0 = x
    if align == "C":
        x0 = x + (width - total) / 2
    pdf.set_text_color(25, 25, 25)
    for seg, it in parts:
        if not seg:
            continue
        pdf.set_font(italic_font if it else reg_font, size=size)
        pdf.set_xy(x0, y)
        pdf.cell(pdf.get_string_width(seg), lh, seg, align="L")
        x0 += pdf.get_string_width(seg)


def docx_title_lines(text, size_pt, max_mm):
    """Estimate Word line wrapping for a centered title (same rule as PDF:
    fill the first line greedily; trailing characters go to a second line)."""
    mm = size_pt * 0.3528

    def cw(ch):
        if ord(ch) > 0x2E7F:      # CJK / full-width
            return mm
        if ch == " ":
            return mm * 0.28
        return mm * 0.52          # Latin letters / digits

    lines, cur, cur_w = [], "", 0.0
    for ch in text:
        w = cw(ch)
        if cur and cur_w + w > max_mm:
            lines.append(cur)
            cur, cur_w = ch, w
        else:
            cur += ch
            cur_w += w
    if cur:
        lines.append(cur)
    return lines


def render_table(pdf, spec, LM, PAGE_W, y_bottom=282):
    columns = spec.get("columns", [])
    rows = spec.get("rows", [])
    ncol = len(columns)
    widths = spec.get("widths") or [PAGE_W / max(ncol, 1)] * max(ncol, 1)
    widths = [min(w, PAGE_W) for w in widths]
    body_size, lh = 7.5, 3.7
    xs = [LM + sum(widths[:i]) for i in range(ncol + 1)]
    x_last = xs[-1]

    def draw_caption(text, font, size, align="C"):
        pdf.set_font(font, size=size)
        pdf.set_text_color(20, 20, 20)
        pdf.multi_cell(PAGE_W, size * 0.53, text, align=align)
        pdf.set_x(LM)
        pdf.ln(0.8)

    def row_lines(texts):
        pdf.set_font("Song", size=body_size)
        return [wrap_text(pdf, render_text(str(t)), widths[i] - 2.4)
                for i, t in enumerate(texts)]

    header = row_lines(columns)
    h_h = max(len(x) for x in header) * lh + 1.8
    body = [row_lines(r) for r in rows]
    r_h = [max(len(x) for x in rl) * lh + 1.6 for rl in body]

    # Keep the whole table on one page when possible.
    cap_est = 12.0  # Chinese + English captions and spacing
    total_h = h_h + sum(r_h) + cap_est
    if pdf.get_y() + total_h > y_bottom:
        pdf.add_page()

    if spec.get("caption"):
        draw_caption(spec["caption"], "Hei", 9)
    if spec.get("caption_en"):
        draw_caption(spec["caption_en"], "Song", 7.5)
    pdf.ln(0.5)

    def draw_header():
        if pdf.get_y() + h_h > y_bottom:
            pdf.add_page()
        y = pdf.get_y()
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.8)
        pdf.line(xs[0], y, x_last, y)          # top thick line
        pdf.set_font("Hei", size=body_size)
        pdf.set_text_color(20, 20, 20)
        for ci, lines in enumerate(header):
            for k, ln in enumerate(lines):
                pdf.set_xy(xs[ci] + 1.2, y + 0.9 + k * lh)
                pdf.cell(widths[ci] - 2.4, lh, ln, align="C")
        hy = y + h_h
        pdf.set_line_width(0.3)
        pdf.line(xs[0], hy, x_last, hy)        # header underline
        pdf.set_xy(LM, hy)

    def draw_body_row(rl, rh):
        if pdf.get_y() + rh > y_bottom:
            pdf.add_page()
            draw_header()
        y = pdf.get_y()
        pdf.set_font("Song", size=body_size)
        pdf.set_text_color(25, 25, 25)
        for ci, lines in enumerate(rl):
            for k, ln in enumerate(lines):
                draw_mixed(pdf, xs[ci] + 1.2, y + 0.8 + k * lh,
                           widths[ci] - 2.4, lh, ln, body_size)
        pdf.set_xy(LM, y + rh)

    draw_header()
    for rl, rh in zip(body, r_h):
        draw_body_row(rl, rh)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.8)
    pdf.line(xs[0], pdf.get_y(), x_last, pdf.get_y())  # bottom thick line
    pdf.ln(2.5)


def render_figure(pdf, spec, LM, PAGE_W, y_bottom=282):
    path = Path(spec["path"])
    if not path.is_absolute():
        path = WORK.parent / path
    if not path.exists():
        print("WARN: figure not found:", path)
        return
    w = spec.get("width", min(PAGE_W * 0.72, 120))
    try:
        iw, ih = png_size(path)
        h = w * ih / iw
    except Exception:  # noqa: BLE001
        h = w * 0.7
    if pdf.get_y() + h > y_bottom - 18:
        pdf.add_page()
    x = LM + (PAGE_W - w) / 2
    pdf.image(str(path), x, pdf.get_y(), w, h)
    pdf.set_xy(LM, pdf.get_y() + h)
    pdf.ln(2)
    if spec.get("caption"):
        pdf.set_font("Hei", size=9)
        pdf.set_text_color(20, 20, 20)
        pdf.multi_cell(PAGE_W, 4.8, spec["caption"], align="C")
        pdf.set_x(LM)
    if spec.get("caption_en"):
        pdf.set_font("TimesR" if EN else "Song", size=7.5)
        pdf.set_text_color(90, 90, 90)
        pdf.multi_cell(PAGE_W, 4.0, spec["caption_en"], align="C")
        pdf.set_x(LM)
    pdf.ln(3)


def build_pdf(ref_keys):
    pdf = ReviewPDF(format="A4")
    pdf.set_margins(22, 20, 22)
    pdf.set_auto_page_break(True, margin=20)
    pdf.add_font("Song", fname=str(SIM_SUN))
    pdf.add_font("Hei", fname=str(SIM_HEI))
    pdf.add_font("TimesR", fname=str(FONT_DIR / "times.ttf"))
    pdf.add_font("TimesB", fname=str(FONT_DIR / "timesbd.ttf"))
    pdf.add_font("TimesI", fname=str(FONT_DIR / "timesi.ttf"))
    pdf.set_title(CONTENT["title"])
    pdf.set_author("Zotero Literature Compilation" if EN else "Zotero 文献整理")
    pdf.alias_nb_pages()
    pdf.add_page()

    LM, RM = pdf.l_margin, pdf.r_margin
    PAGE_W = pdf.w - LM - RM
    # Keep blocks inside the auto page-break trigger (297 - 20 = 277 mm) with a
    # small buffer, so paragraphs never spill a single line onto the next page
    # and content never overlaps the footer.
    Y_BOTTOM = 272

    def set_font(name, size, color=(20, 20, 20)):
        pdf.set_font(name, size=size)
        pdf.set_text_color(*color)

    # ---- Title (SimHei 22 pt, centered; balanced character wrapping) ----
    set_font("TimesB" if EN else "Hei", 22)
    for ln in wrap_title(pdf, CONTENT["title"], PAGE_W):
        pdf.multi_cell(0, 11.5, ln, align="C")
        pdf.set_x(LM)
    pdf.ln(1.5)
    set_font("TimesR" if EN else "Song", 9, (110, 116, 125))
    pdf.multi_cell(0, 4.8, CONTENT["subtitle"], align="C")
    pdf.set_x(LM)
    pdf.ln(3)

    # ---- 摘要 / 关键词 / Abstract / Keywords（标签与内容同行、续行顶格）----
    def abs_block(label, text, label_font, label_size, body_size,
                  block_mm, first_mm):
        lh = 5.6
        y = pdf.get_y()
        x_first = LM + block_mm + first_mm
        pdf.set_font(label_font, size=label_size)
        lw = pdf.get_string_width(label)
        pdf.set_font("TimesR" if EN else "Song", size=body_size)
        width_first = LM + PAGE_W - x_first - lw
        lines = wrap_text(pdf, text, width_first)
        total = lh * len(lines)
        if y + total > Y_BOTTOM:
            pdf.add_page()
            y = pdf.get_y()
            x_first = LM + block_mm + first_mm
        pdf.set_font(label_font, size=label_size)
        pdf.set_text_color(25, 25, 25)
        pdf.set_xy(x_first, y)
        pdf.cell(lw, lh, label)
        pdf.set_font("TimesR" if EN else "Song", size=body_size)
        pdf.set_text_color(35, 35, 35)
        draw_mixed(pdf, x_first + lw, y, width_first, lh, lines[0], body_size)
        for k, ln in enumerate(lines[1:], start=1):
            draw_mixed(pdf, LM + block_mm, y + k * lh, PAGE_W - block_mm,
                       lh, ln, body_size)
        pdf.set_xy(LM, y + total)
        pdf.ln(1.2)

    if CONTENT.get("intro"):
        abs_block("Abstract: " if EN else "摘  要：", pdf_safe(render_text(CONTENT["intro"])), "TimesB" if EN else "Hei", 9, 9, 0, 0)
    if CONTENT.get("keywords"):
        abs_block("Keywords: " if EN else "关键词：", CONTENT["keywords"], "TimesB" if EN else "Hei", 9, 9, 0, 0)
    if not EN and CONTENT.get("title_en"):
        pdf.ln(2)
        pdf.set_font("TimesB", size=14)
        pdf.set_text_color(15, 15, 15)
        pdf.multi_cell(PAGE_W, 7.5, CONTENT["title_en"], align="L")
        pdf.set_x(LM)
        pdf.ln(2)
    if not EN and CONTENT.get("abstract_en"):
        abs_block("Abstract:", CONTENT["abstract_en"], "TimesB", 10.5, 10.5, 0, 0)
    if not EN and CONTENT.get("keywords_en"):
        abs_block("Keywords:", CONTENT["keywords_en"], "TimesB", 10.5, 10.5, 0, 0)
    pdf.ln(1)

    def section(title, sec_no):
        pdf.ln(2.5)
        set_font("TimesB" if EN else "Song", 14, (15, 15, 15))
        label = f"{sec_no}  {title}" if sec_no else title
        pdf.multi_cell(0, 7.6, label, align="L")
        pdf.set_x(LM)
        pdf.ln(1.8)

    def sub(title, sec_no, sub_no):
        pdf.ln(1.2)
        set_font("TimesB" if EN else "Hei", 10.5, (20, 20, 20))
        pdf.multi_cell(0, 5.6, f"{sec_no}.{sub_no}  {title}", align="L")
        pdf.set_x(LM)
        pdf.ln(0.8)

    def item(text, bullet=False):
        lh = 5.6
        indent = 7.4  # two full-width characters at 10.5 pt
        set_font("TimesR" if EN else "Song", 10.5, (25, 25, 25))
        lines = wrap_text(pdf, text, PAGE_W - indent)
        y = pdf.get_y()
        if y + len(lines) * lh > Y_BOTTOM:
            pdf.add_page()
            y = pdf.get_y()
        draw_mixed(pdf, LM + indent, y, PAGE_W - indent, lh, lines[0], 10.5)
        for k, ln in enumerate(lines[1:], start=1):
            draw_mixed(pdf, LM, y + k * lh, PAGE_W, lh, ln, 10.5)
        pdf.set_xy(LM, y + len(lines) * lh)
        pdf.ln(0.5)

    sec_idx = 0
    for sec in CONTENT["sections"]:
        n, text = section_number(sec["heading"])
        sec_idx = n if n else sec_idx + 1
        section(text, sec_idx)
        sub_idx = 0
        for group in sec.get("groups", []):
            if group.get("table"):
                render_table(pdf, group["table"], LM, PAGE_W, Y_BOTTOM)
            if group.get("image"):
                render_figure(pdf, group["image"], LM, PAGE_W, Y_BOTTOM)
            if group.get("sub"):
                sn, stext = sub_number(group["sub"])
                sub_idx = sn if sn else sub_idx + 1
                sub(stext, sec_idx, sub_idx)
            for text in group.get("items", []):
                item(pdf_safe(render_text(text)))

    # ---- References (SimSun 7.5 pt, alphabetical) ----
    section("References" if EN else "参考文献", "")
    set_font("TimesR" if EN else "Song", 7.5, (25, 25, 25))
    hang = 5.3   # two full-width characters at 7.5 pt
    lh = 3.9
    for key in sorted(ref_keys, key=ref_sort_key):
        txt = pdf_safe(ref_text(key))
        if not txt:
            continue
        lines = wrap_text(pdf, txt, PAGE_W - hang)
        y = pdf.get_y()
        if y + len(lines) * lh > Y_BOTTOM:
            pdf.add_page()
            y = pdf.get_y()
        draw_mixed(pdf, LM, y, PAGE_W - hang, lh, lines[0], 7.5)
        for k, ln in enumerate(lines[1:], start=1):
            draw_mixed(pdf, LM + hang, y + k * lh, PAGE_W - hang, lh, ln, 7.5)
        pdf.set_xy(LM, y + len(lines) * lh)
        pdf.ln(0.9)

    out = OUT_DIR / (OUT_NAME + ".pdf")
    pdf.output(str(out))
    print(f"wrote {out} ({out.stat().st_size} bytes, {pdf.pages_count} pages)")
    return out


# --------------------------------------------------------------------------
# Verification
# --------------------------------------------------------------------------

def verify(docx_path):
    check = Document(docx_path)
    texts = [p.text for p in check.paragraphs]
    print("docx reopened: paragraphs =", len(texts))
    suspicious = set("\u9225\u604A\u94FF\u4E67\u788C\u4E6A\u4FF9\u50E3\u614F\u8133\u923C\u4E76\u617A\u4E7A\uE6A9\u20AC")
    bad_pat = re.compile("[\uE000-\uF8FF\uFB00-\uFB06]")
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    hits = {ch for ch in xml if ch in suspicious or bad_pat.match(ch)}
    print("docx garble scan:", "clean" if not hits else "HAS GARBLED " + ",".join("U+%04X" % ord(h) for h in sorted(hits)))
    return texts


def main():
    order = build_citations(CONTENT["sections"])
    print("citations found:", len(order))
    validate_citations(order)
    docx = build_docx(order)
    build_pdf(order)
    verify(docx)
    import fitz  # noqa: E402

    pdf_path = OUT_DIR / (OUT_NAME + ".pdf")
    doc = fitz.open(pdf_path)
    text = "\n".join(pg.get_text() for pg in doc)
    suspicious = set("\u9225\u604A\u94FF\u4E67\u788C\u4E6A\u4FF9\u50E3\u614F\u8133\u923C\u4E76\u617A\u4E7A\uE6A9\u20AC")
    hits = {ch for ch in text if ch in suspicious}
    print(f"pdf pages: {doc.page_count} | garble: {'clean' if not hits else sorted(hits)}")
    first_lines = [ln for ln in doc[0].get_text().splitlines() if ln.strip()][:3]
    print("pdf first lines:", first_lines)


if __name__ == "__main__":
    main()
