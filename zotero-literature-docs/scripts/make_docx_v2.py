import json
import re
import sys
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:  # noqa: BLE001
    pass
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
sys.path.insert(0, str(Path(__file__).parent / "pylibs2"))

import make_papers_doc as m  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

WORK = Path(__file__).parent
BIB = WORK.parent / "outputs" / "genetic_transformation.bib"
OUT = WORK.parent / "outputs" / "基因遗传转化文献.docx"
(WORK.parent / "outputs").mkdir(exist_ok=True)
CONFIG = WORK / "doc_config.json"

DARK = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)
LIGHT = RGBColor(0x6E, 0x6E, 0x6E)
BLACK = RGBColor(0x14, 0x14, 0x14)


def set_east_asia(run, name="Microsoft YaHei"):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def para(doc, size, color=BLACK, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    if align:
        p.alignment = align
    r = p.add_run()
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    set_east_asia(r)
    return p, r


def load_config():
    default_sections = [
        {"title": "一、综述与方法类", "keys": None},
        {"title": "二、兰花（蝴蝶兰等）转化研究", "keys": None},
    ]
    if CONFIG.exists():
        cfg = json.loads(CONFIG.read_text(encoding="utf-8"))
        cfg.setdefault("title", "文献清单")
        cfg.setdefault("output", cfg["title"])
        cfg.setdefault("sections", default_sections)
        return cfg
    return {"title": "基因遗传转化相关文献", "output": "基因遗传转化文献.docx", "sections": default_sections}


def split_sections(entries, sections):
    """Return list of (title, items). Honors per-section 'keys' when present."""
    if all(s.get("keys") for s in sections):
        result = []
        for s in sections:
            wanted = set(s["keys"])
            result.append((s["title"], [e for e in entries if e["key"] in wanted]))
        return result
    half = len(entries) // 2
    return [(sections[0]["title"], entries[:half]), (sections[1]["title"], entries[half:])]


def main():
    entries = m.load_entries()
    cfg = load_config()
    sections = split_sections(entries, cfg["sections"])
    out = WORK.parent / "outputs" / (cfg["output"] + ".docx")
    total = len(entries)
    assert total > 0

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.line_spacing = 1.2

    # Title
    p, r = para(doc, 20, DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    r.text = cfg["title"]
    p, r = para(doc, 10, LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    r.text = f"来源：Zotero 个人文献库    |    整理日期：2026-08-04    |    共 {total} 篇"

    def section(title):
        p, r = para(doc, 14, DARK, bold=True, after=6)
        p.paragraph_format.space_before = Pt(10)
        r.text = title

    def entry(idx, e):
        title = m.clean(e["title"])
        authors = e["author"]
        journal_bits = [m.clean(e["journal"])]
        if e["volume"]:
            vol = f'{m.clean(e["volume"])}({m.clean(e["number"])})' if e["number"] else m.clean(e["volume"])
            journal_bits.append(vol)
        if e["pages"]:
            journal_bits.append(f"pp. {m.clean(e['pages'])}")
        journal_bits.append(f"({m.clean(e['year'])})")
        meta = ", ".join(x for x in journal_bits if x)
        if e["doi"]:
            meta += f'  DOI: https://doi.org/{m.clean(e["doi"])}'

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(f"{idx}. {title}")
        r1.bold = True
        r1.font.size = Pt(11)
        set_east_asia(r1)
        r2 = p.add_run()
        r2.add_break(WD_BREAK.LINE)
        r2.text = authors
        r2.font.size = Pt(10)
        set_east_asia(r2)
        r3 = p.add_run()
        r3.add_break(WD_BREAK.LINE)
        r3.text = meta
        r3.font.size = Pt(10)
        r3.font.color.rgb = GRAY
        set_east_asia(r3)

        if e["abstract"]:
            ab = m.clean(e["abstract"])
            if len(ab) > 380:
                ab = ab[:380].rsplit(" ", 1)[0] + " ..."
            pa, ra = para(doc, 9, LIGHT, after=8)
            ra.text = "摘要：" + ab

    idx = 0
    for title, items in sections:
        section(title)
        for e in items:
            idx += 1
            entry(idx, e)

    doc.save(out)
    print(f"wrote {out} ({out.stat().st_size} bytes)")

    # Round-trip validation
    check = Document(out)
    texts = [p.text for p in check.paragraphs]
    numbered = [t for t in texts if t and t[0].isdigit() and ". " in t[:4]]
    print("reopened OK: paragraphs =", len(texts), "| numbered entries =", len(numbered))

    # Garble scan
    suspicious = set("\u9225\u604A\u94FF\u4E67\u788C\u4E6A\u4FF9\u50E3\u614F\u8133\u923C\u4E76\u617A\u4E7A\uE6A9\u20AC")
    bad_pat = re.compile("[\uE000-\uF8FF\uFB00-\uFB06]")
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    hits = {ch for ch in xml if ch in suspicious or bad_pat.match(ch)}
    print("garble scan:", "clean" if not hits else "HAS GARBLED " + ",".join("U+%04X" % ord(h) for h in sorted(hits)))


if __name__ == "__main__":
    main()
